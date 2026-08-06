import csv
import json
import re
import uuid
from dataclasses import dataclass
from datetime import UTC, datetime
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any

from azure.core.exceptions import ResourceNotFoundError
from app.azure_credential import get_azure_credential
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import (
    HnswAlgorithmConfiguration,
    SearchableField,
    SearchField,
    SearchFieldDataType,
    SearchIndex,
    SimpleField,
    VectorSearch,
    VectorSearchProfile,
)
from azure.search.documents.models import VectorizedQuery
from azure.storage.blob import BlobServiceClient, ContentSettings

from app.foundry_client import create_embeddings, load_settings
from app.config import env_int, first_env
from app.security import UserScope

MAX_DOCUMENT_BYTES = 12 * 1024 * 1024
CHUNK_TARGET_CHARS = 1400
CHUNK_OVERLAP_CHARS = 180
DEFAULT_TOP_K = 6
SEARCH_SCHEMA_VERSION = "v2"
VECTOR_PROFILE_NAME = "foundry-rag-vector-profile"
VECTOR_ALGORITHM_NAME = "foundry-rag-hnsw"

SUPPORTED_EXTENSIONS = {
    ".csv",
    ".docx",
    ".html",
    ".htm",
    ".json",
    ".log",
    ".md",
    ".pdf",
    ".text",
    ".txt",
    ".xml",
}


@dataclass(frozen=True)
class RagSearchSettings:
    endpoint: str | None
    index_name: str
    embedding_model: str
    embedding_dimensions: int | None
    storage_account_url: str | None
    storage_container_name: str

    @property
    def is_configured(self) -> bool:
        return bool(self.endpoint and self.index_name and self.storage_account_url and self.storage_container_name)


@dataclass(frozen=True)
class UploadedDocument:
    id: str
    filename: str
    content_type: str | None
    byte_size: int
    chunk_count: int
    blob_name: str | None
    blob_url: str | None
    created_at: str


@dataclass(frozen=True)
class RetrievedChunk:
    document_id: str
    filename: str
    chunk_index: int
    content: str
    score: float
    blob_url: str | None


def load_rag_search_settings() -> RagSearchSettings:
    foundry_settings = load_settings()
    return RagSearchSettings(
        endpoint=first_env("AZURE_SEARCH_ENDPOINT", "FOUNDRY_SEARCH_ENDPOINT"),
        index_name=f"{first_env('AZURE_SEARCH_INDEX_NAME', 'FOUNDRY_SEARCH_INDEX_NAME', default='foundry-document-rag')}-{SEARCH_SCHEMA_VERSION}",
        embedding_model=foundry_settings.embedding_model,
        embedding_dimensions=env_int(
            "FOUNDRY_EMBEDDING_DIMENSIONS",
            0,
            minimum=0,
        ) or None,
        storage_account_url=first_env(
            "AZURE_STORAGE_ACCOUNT_URL",
            "FOUNDRY_STORAGE_ACCOUNT_URL",
        ),
        storage_container_name=first_env(
            "AZURE_STORAGE_CONTAINER_NAME",
            "FOUNDRY_STORAGE_CONTAINER_NAME",
            default="foundry-rag-documents",
        ) or "foundry-rag-documents",
    )


def add_document(
    *,
    scope: UserScope,
    filename: str,
    content_type: str | None,
    data: bytes,
) -> dict[str, Any]:
    settings = _require_search_settings()
    safe_filename = Path(filename or "uploaded-document").name
    if not data:
        raise ValueError("Uploaded document is empty.")
    if len(data) > MAX_DOCUMENT_BYTES:
        raise ValueError("Uploaded document exceeds the 12 MB limit.")
    if not _is_supported_file(safe_filename, content_type):
        raise ValueError(
            "Unsupported document type. Upload PDF, DOCX, TXT, Markdown, CSV, JSON, HTML, XML, or log files."
        )

    extracted_text = _extract_text(safe_filename, content_type, data)
    if not extracted_text.strip():
        raise ValueError("No readable text was found in the uploaded document.")

    chunks = _chunk_text(extracted_text)
    if not chunks:
        raise ValueError("No readable text chunks were found in the uploaded document.")

    document_id = str(uuid.uuid4())
    created_at = datetime.now(UTC).isoformat()
    blob_name = f"documents/{scope.tenant_id}/{scope.user_id}/{document_id}/{safe_filename}"
    blob_url = _upload_original_document(
        settings=settings,
        blob_name=blob_name,
        filename=safe_filename,
        content_type=content_type,
        data=data,
    )

    embeddings = _embed_text_batches(chunks, settings.embedding_model)
    vectors = embeddings["vectors"]
    embedding_dimensions = settings.embedding_dimensions or (len(vectors[0]) if vectors else None)
    if not embedding_dimensions:
        raise RuntimeError("Foundry embedding response did not include vector dimensions.")
    _ensure_search_index(settings, embedding_dimensions)

    search_documents = [
        {
            "id": f"{document_id}-{index}",
            "document_id": document_id,
            "tenant_id": scope.tenant_id,
            "owner_id": scope.user_id,
            "filename": safe_filename,
            "content_type": content_type or "",
            "byte_size": len(data),
            "blob_name": blob_name,
            "blob_url": blob_url,
            "chunk_index": index,
            "content": chunk,
            "content_vector": vector,
            "created_at": created_at,
        }
        for index, (chunk, vector) in enumerate(zip(chunks, vectors, strict=True), start=1)
    ]
    with _create_search_client(settings) as search_client:
        search_client.upload_documents(search_documents)

    return {
        "document": document_to_dict(
            UploadedDocument(
                id=document_id,
                filename=safe_filename,
                content_type=content_type,
                byte_size=len(data),
                chunk_count=len(chunks),
                blob_name=blob_name,
                blob_url=blob_url,
                created_at=created_at,
            )
        ),
        "embedding": {
            "model": embeddings["model"],
            "duration_ms": embeddings["duration_ms"],
            "dimensions": embedding_dimensions,
            "foundry_request": embeddings["foundry_request"],
            "foundry_response": embeddings["foundry_response"],
        },
    }


def list_documents(scope: UserScope) -> list[UploadedDocument]:
    settings = _require_search_settings()
    documents: dict[str, UploadedDocument] = {}
    with _create_search_client(settings) as search_client:
        try:
            results = search_client.search(
                search_text="*",
                filter=_owner_filter(scope),
                select=[
                    "document_id",
                    "filename",
                    "content_type",
                    "byte_size",
                    "blob_name",
                    "blob_url",
                    "chunk_index",
                    "created_at",
                ],
                top=1000,
            )
            for result in results:
                document_id = str(result["document_id"])
                existing = documents.get(document_id)
                chunk_count = (existing.chunk_count if existing else 0) + 1
                documents[document_id] = UploadedDocument(
                    id=document_id,
                    filename=str(result["filename"]),
                    content_type=str(result.get("content_type") or "") or None,
                    byte_size=int(result.get("byte_size") or 0),
                    chunk_count=chunk_count,
                    blob_name=str(result.get("blob_name") or "") or None,
                    blob_url=str(result.get("blob_url") or "") or None,
                    created_at=str(result["created_at"]),
                )
        except ResourceNotFoundError:
            return []
    return sorted(documents.values(), key=lambda document: document.created_at, reverse=True)


def delete_document(scope: UserScope, document_id: str) -> bool:
    settings = _require_search_settings()
    escaped_document_id = _escape_search_filter_value(document_id)
    with _create_search_client(settings) as search_client:
        chunk_keys = []
        blob_names: set[str] = set()
        try:
            results = search_client.search(
                search_text="*",
                filter=f"{_owner_filter(scope)} and document_id eq '{escaped_document_id}'",
                select=["id", "blob_name"],
                top=1000,
            )
            for result in results:
                chunk_keys.append({"id": str(result["id"])})
                blob_name = str(result.get("blob_name") or "")
                if blob_name:
                    blob_names.add(blob_name)
        except ResourceNotFoundError:
            return False
        if not chunk_keys:
            return False
        search_client.delete_documents(documents=chunk_keys)
    for blob_name in blob_names:
        _delete_original_document(settings, blob_name)
    return True


def retrieve_document_chunks(
    scope: UserScope,
    query: str,
    *,
    limit: int = DEFAULT_TOP_K,
) -> dict[str, Any]:
    settings = _require_search_settings()
    embedding = create_embeddings(inputs=[query], model=settings.embedding_model)
    query_vector = embedding["vectors"][0]
    vector_query = VectorizedQuery(
        vector=query_vector,
        k_nearest_neighbors=limit,
        fields="content_vector",
    )
    with _create_search_client(settings) as search_client:
        try:
            results = search_client.search(
                search_text=query,
                filter=_owner_filter(scope),
                vector_queries=[vector_query],
                select=["document_id", "filename", "chunk_index", "content", "blob_url"],
                top=limit,
            )
            chunks = [
                RetrievedChunk(
                    document_id=str(result["document_id"]),
                    filename=str(result["filename"]),
                    chunk_index=int(result["chunk_index"]),
                    content=str(result["content"]),
                    score=float(result.get("@search.score") or 0),
                    blob_url=str(result.get("blob_url") or "") or None,
                )
                for result in results
            ]
        except ResourceNotFoundError:
            chunks = []
    return {
        "chunks": chunks,
        "embedding": {
            "model": embedding["model"],
            "duration_ms": embedding["duration_ms"],
            "dimensions": len(query_vector),
            "foundry_request": embedding["foundry_request"],
            "foundry_response": embedding["foundry_response"],
        },
    }


def document_to_dict(document: UploadedDocument) -> dict[str, Any]:
    return {
        "id": document.id,
        "filename": document.filename,
        "content_type": document.content_type,
        "byte_size": document.byte_size,
        "chunk_count": document.chunk_count,
        "blob_name": document.blob_name,
        "blob_url": document.blob_url,
        "created_at": document.created_at,
    }


def chunk_to_dict(chunk: RetrievedChunk) -> dict[str, Any]:
    return {
        "document_id": chunk.document_id,
        "filename": chunk.filename,
        "chunk_index": chunk.chunk_index,
        "content": chunk.content,
        "score": chunk.score,
        "blob_url": chunk.blob_url,
    }


def build_grounded_prompt(question: str, chunks: list[RetrievedChunk]) -> str:
    sources = "\n\n".join(
        f"[{index}] {chunk.filename} (chunk {chunk.chunk_index})\n{chunk.content}"
        for index, chunk in enumerate(chunks, start=1)
    )
    return (
        "Answer the user's question using only the document excerpts below. "
        "Cite supporting excerpts inline with bracketed source numbers like [1]. "
        "If the excerpts do not contain the answer, say you could not find it in the uploaded documents.\n\n"
        f"Document excerpts:\n{sources or 'No relevant excerpts were retrieved.'}\n\n"
        f"Question: {question}"
    )


def build_rag_system_prompt(system_prompt: str) -> str:
    base_prompt = system_prompt.strip() or "You are a concise, helpful assistant."
    return (
        f"{base_prompt}\n\n"
        "You are answering a document-grounded RAG question. Do not use outside knowledge for facts "
        "about the uploaded documents, and cite the provided excerpts with [source number] references."
    )


def _require_search_settings() -> RagSearchSettings:
    settings = load_rag_search_settings()
    if not settings.is_configured:
        raise RuntimeError(
            "Document Q&A is not configured. Set AZURE_SEARCH_ENDPOINT, AZURE_SEARCH_INDEX_NAME, "
            "AZURE_STORAGE_ACCOUNT_URL, and AZURE_STORAGE_CONTAINER_NAME."
        )
    return settings


def _create_search_client(settings: RagSearchSettings) -> SearchClient:
    return SearchClient(
        endpoint=settings.endpoint or "",
        index_name=settings.index_name,
        credential=get_azure_credential(),
    )


def _ensure_search_index(settings: RagSearchSettings, embedding_dimensions: int) -> None:
    index_client = SearchIndexClient(
        endpoint=settings.endpoint or "",
        credential=get_azure_credential(),
    )
    fields = [
        SimpleField(name="id", type=SearchFieldDataType.String, key=True, filterable=True),
        SimpleField(name="document_id", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="tenant_id", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="owner_id", type=SearchFieldDataType.String, filterable=True),
        SearchableField(name="filename", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="content_type", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="byte_size", type=SearchFieldDataType.Int64, filterable=True),
        SimpleField(name="blob_name", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="blob_url", type=SearchFieldDataType.String),
        SimpleField(name="chunk_index", type=SearchFieldDataType.Int32, filterable=True, sortable=True),
        SearchableField(name="content", type=SearchFieldDataType.String),
        SearchField(
            name="content_vector",
            type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
            searchable=True,
            vector_search_dimensions=embedding_dimensions,
            vector_search_profile_name=VECTOR_PROFILE_NAME,
        ),
        SimpleField(name="created_at", type=SearchFieldDataType.String, filterable=True, sortable=True),
    ]
    index = SearchIndex(
        name=settings.index_name,
        fields=fields,
        vector_search=VectorSearch(
            algorithms=[HnswAlgorithmConfiguration(name=VECTOR_ALGORITHM_NAME)],
            profiles=[
                VectorSearchProfile(
                    name=VECTOR_PROFILE_NAME,
                    algorithm_configuration_name=VECTOR_ALGORITHM_NAME,
                )
            ],
        ),
    )
    index_client.create_or_update_index(index)
    index_client.close()


def _upload_original_document(
    *,
    settings: RagSearchSettings,
    blob_name: str,
    filename: str,
    content_type: str | None,
    data: bytes,
) -> str:
    blob_service_client = BlobServiceClient(
        account_url=settings.storage_account_url or "",
        credential=get_azure_credential(),
    )
    try:
        container_client = blob_service_client.get_container_client(settings.storage_container_name)
        if not container_client.exists():
            container_client.create_container()
        blob_client = container_client.get_blob_client(blob_name)
        blob_client.upload_blob(
            data,
            overwrite=True,
            content_settings=ContentSettings(
                content_type=content_type or "application/octet-stream",
                content_disposition=f'attachment; filename="{filename.replace(chr(34), "")}"',
            ),
        )
        return blob_client.url
    finally:
        blob_service_client.close()


def _delete_original_document(settings: RagSearchSettings, blob_name: str) -> None:
    blob_service_client = BlobServiceClient(
        account_url=settings.storage_account_url or "",
        credential=get_azure_credential(),
    )
    blob_client = blob_service_client.get_blob_client(
        container=settings.storage_container_name,
        blob=blob_name,
    )
    try:
        blob_client.delete_blob(delete_snapshots="include")
    except ResourceNotFoundError:
        return
    finally:
        blob_service_client.close()


def _embed_text_batches(chunks: list[str], model: str) -> dict[str, Any]:
    vectors: list[list[float]] = []
    total_duration_ms = 0
    first_request: dict[str, Any] | None = None
    last_response: dict[str, Any] | None = None
    for index in range(0, len(chunks), 16):
        batch = chunks[index : index + 16]
        embedding = create_embeddings(inputs=batch, model=model)
        vectors.extend(embedding["vectors"])
        total_duration_ms += embedding["duration_ms"]
        first_request = first_request or embedding["foundry_request"]
        last_response = embedding["foundry_response"]
    return {
        "model": model,
        "vectors": vectors,
        "duration_ms": total_duration_ms,
        "foundry_request": first_request,
        "foundry_response": last_response,
    }


def _is_supported_file(filename: str, content_type: str | None) -> bool:
    extension = Path(filename).suffix.lower()
    if extension in SUPPORTED_EXTENSIONS:
        return True
    return bool(content_type and (content_type.startswith("text/") or content_type == "application/json"))


def _extract_text(filename: str, content_type: str | None, data: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        return _extract_pdf_text(data)
    if extension == ".docx":
        return _extract_docx_text(data)
    if extension == ".csv":
        return _extract_csv_text(data)
    if extension == ".json" or content_type == "application/json":
        return _extract_json_text(data)
    return _decode_text(data)


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("PDF upload support requires the pypdf package.") from exc

    reader = PdfReader(BytesIO(data))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("DOCX upload support requires the python-docx package.") from exc

    document = Document(BytesIO(data))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_rows: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
    return "\n\n".join([*paragraphs, *table_rows])


def _extract_csv_text(data: bytes) -> str:
    text = _decode_text(data)
    rows = csv.reader(StringIO(text))
    return "\n".join(" | ".join(cell.strip() for cell in row if cell.strip()) for row in rows)


def _extract_json_text(data: bytes) -> str:
    text = _decode_text(data)
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        return text
    return json.dumps(parsed, indent=2, ensure_ascii=False)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _chunk_text(text: str) -> list[str]:
    normalized = re.sub(r"\r\n?", "\n", text)
    paragraphs = [paragraph.strip() for paragraph in re.split(r"\n{2,}", normalized) if paragraph.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(paragraph) > CHUNK_TARGET_CHARS:
            if current:
                chunks.append(current.strip())
                current = ""
            chunks.extend(_split_long_text(paragraph))
            continue
        separator = "\n\n" if current else ""
        if len(current) + len(separator) + len(paragraph) <= CHUNK_TARGET_CHARS:
            current = f"{current}{separator}{paragraph}"
        else:
            if current:
                chunks.append(current.strip())
            current = paragraph
    if current:
        chunks.append(current.strip())
    return [chunk for chunk in chunks if chunk.strip()]


def _split_long_text(text: str) -> list[str]:
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + CHUNK_TARGET_CHARS, len(text))
        chunks.append(text[start:end].strip())
        if end == len(text):
            break
        start = max(end - CHUNK_OVERLAP_CHARS, start + 1)
    return [chunk for chunk in chunks if chunk]


def _escape_search_filter_value(value: str) -> str:
    return value.replace("'", "''")


def _owner_filter(scope: UserScope) -> str:
    tenant_id = _escape_search_filter_value(scope.tenant_id)
    owner_id = _escape_search_filter_value(scope.user_id)
    return f"tenant_id eq '{tenant_id}' and owner_id eq '{owner_id}'"
